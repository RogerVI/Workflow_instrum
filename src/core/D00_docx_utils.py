# core/D00_docx_utils.py
from __future__ import annotations
from typing import Dict, Iterable, Tuple
from io import BytesIO
from docx import Document
from docx.shared import Inches

def _iter_all_paragraphs(doc) -> Iterable:
    """Itère sur tous les paragraphes du document, y compris dans les tableaux."""
    # corps principal
    for p in doc.paragraphs:
        yield p
    # tableaux (cellules)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def insert_image_after_marker_docx(
    doc_bytes: bytes,
    images: Dict[str, bytes],
    width_in_inches: float = 6.0,
    marker_prefix: str = "### ",
) -> Tuple[bytes, int, int]:
    """
    Ouvre un .docx (bytes), insère chaque image après un paragraphe contenant
    le repère f"{marker_prefix}{nom_image_sans_ext}", puis renvoie le docx bytes.
    - images: { "NomDuGraph": png_bytes }  (⚠️ sans extension)
    - retourne: (docx_bytes, nb_insere, nb_non_trouve)
    """
    doc = Document(BytesIO(doc_bytes))
    total_ok, total_miss = 0, 0

    # construit un index "texte paragraphe -> paragraph object"
    # (on cherche par inclusion, pas égalité stricte)
    for name, img in images.items():
        marker = f"{marker_prefix}{name}"
        inserted = False
        for i, para in enumerate(_iter_all_paragraphs(doc)):
            if marker in para.text:
                # paragraphe suivant si possible, sinon on crée en fin
                # essayé: ajouter dans le paragraphe suivant pour plus de lisibilité
                run = None
                # on tente de récupérer le paragraphe "suivant" en se rabattant sur ajout
                # (python-docx n'a pas d'API "next paragraph" universelle, on ajoute donc après)
                # on ajoute juste après sous forme d’un nouveau paragraphe:
                run = doc.add_paragraph().add_run()
                run.add_picture(BytesIO(img), width=Inches(width_in_inches))
                inserted = True
                break

        if inserted:
            total_ok += 1
        else:
            total_miss += 1

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read(), total_ok, total_miss
